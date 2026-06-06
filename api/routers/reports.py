"""
routers/reports.py
Intelligent analytical report generator — PDF, Excel, and editable DOCX.

Key improvements over v1:
 - Narrative prose sections (situation, commodity analysis, forecast outlook, indicators)
 - Multi-commodity support (select several; narrative covers all)
 - Forecast section from the /api/forecast endpoint data
 - FSI / PSI / Early-warning indicators section
 - DOCX export (editable Word document)
 - Fixed: timestamp in Last Report, FSI 0.0/100 labelling, invalid month-31 dates
 - Fixed: district cover page stats, hardcoded "3 Regions" on cover
"""

from fastapi import APIRouter, Query
from fastapi.responses import Response, JSONResponse
from api.database import get_pool
from datetime import datetime, date as date_type, timezone
from io import BytesIO
from typing import List, Optional
import textwrap

# ── ReportLab ─────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak,
)
from reportlab.platypus.flowables import Flowable

# ── python-docx ───────────────────────────────────────────────────────────────
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

router = APIRouter(prefix="/api/reports", tags=["Reports"])

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY        = colors.HexColor("#1B3A6B")
NAVY_LIGHT  = colors.HexColor("#2A4F8F")
GOLD        = colors.HexColor("#F5C842")
RED         = colors.HexColor("#B71C1C")
RED_LIGHT   = colors.HexColor("#EF5350")
ORANGE      = colors.HexColor("#E65100")
AMBER       = colors.HexColor("#F9A825")
GREEN       = colors.HexColor("#2E7D32")
BLUE_LIGHT  = colors.HexColor("#A8C4E8")
GREY_BG     = colors.HexColor("#F8F9FA")
GREY_LINE   = colors.HexColor("#E0E0E0")
GREY_MID    = colors.HexColor("#666666")
GREY_DARK   = colors.HexColor("#333333")
ORANGE_BG   = colors.HexColor("#FFF3E0")
BLUE_BG     = colors.HexColor("#EEF2FF")
GREEN_BG    = colors.HexColor("#F1F8E9")
WHITE       = colors.white
W, H        = A4


# ═══════════════════════════════════════════════════════════════════════════════
# Custom ReportLab flowables
# ═══════════════════════════════════════════════════════════════════════════════

class ColorRect(Flowable):
    def __init__(self, width, height, fill_color):
        Flowable.__init__(self)
        self.width = width; self.height = height; self.fill_color = fill_color
    def draw(self):
        self.canv.setFillColor(self.fill_color)
        self.canv.rect(0, 0, self.width, self.height, fill=1, stroke=0)


class LeftBorderBox(Flowable):
    def __init__(self, content_para, width, bg_color, border_color, border_width=4):
        Flowable.__init__(self)
        self.content_para = content_para; self.width = width
        self.bg_color = bg_color; self.border_color = border_color
        self.border_width = border_width; self._height = None

    def wrap(self, availWidth, availHeight):
        _, h = self.content_para.wrap(self.width - self.border_width - 16, availHeight)
        self._height = h + 20
        return self.width, self._height

    def draw(self):
        h = self._height; c = self.canv
        c.setFillColor(self.bg_color)
        c.rect(0, 0, self.width, h, fill=1, stroke=0)
        c.setFillColor(self.border_color)
        c.rect(0, 0, self.border_width, h, fill=1, stroke=0)
        c.setStrokeColor(self.border_color); c.setLineWidth(0.5)
        c.rect(0, 0, self.width, h, fill=0, stroke=1)
        self.content_para.drawOn(c, self.border_width + 10, 10)


class NarrativeBox(Flowable):
    """Full-width narrative text box with optional heading stripe."""
    def __init__(self, heading, paragraphs, width, bg=None, heading_bg=None):
        Flowable.__init__(self)
        self.heading = heading
        self.paragraphs = paragraphs  # list of Paragraph objects
        self.width = width
        self.bg = bg or GREY_BG
        self.heading_bg = heading_bg or NAVY
        self._height = None
        self._para_heights = []

    def wrap(self, availWidth, availHeight):
        inner_w = self.width - 24
        total_h = 32  # heading stripe
        heights = []
        for p in self.paragraphs:
            _, ph = p.wrap(inner_w, 9999)
            heights.append(ph)
            total_h += ph + 8
        total_h += 16  # bottom padding
        self._height = total_h
        self._para_heights = heights
        return self.width, self._height

    def draw(self):
        c = self.canv; w = self.width; h = self._height
        # background
        c.setFillColor(self.bg)
        c.rect(0, 0, w, h, fill=1, stroke=0)
        # heading stripe
        c.setFillColor(self.heading_bg)
        c.rect(0, h - 32, w, 32, fill=1, stroke=0)
        # border
        c.setStrokeColor(GREY_LINE); c.setLineWidth(0.5)
        c.rect(0, 0, w, h, fill=0, stroke=1)
        # heading text
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(12, h - 20, self.heading.upper())
        # paragraph content
        y = h - 32 - 8
        for p, ph in zip(self.paragraphs, self._para_heights):
            y -= ph
            p.drawOn(c, 12, y)
            y -= 8


# ═══════════════════════════════════════════════════════════════════════════════
# Style helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _style(name, **kwargs):
    defaults = dict(fontName="Helvetica", fontSize=9,
                    textColor=GREY_DARK, spaceAfter=0, spaceBefore=0, leading=13)
    defaults.update(kwargs)
    return ParagraphStyle(name, **defaults)

def _p(text, style):
    return Paragraph(text, style)

def risk_color(score):
    if score >= 277: return RED
    if score >= 199: return RED_LIGHT
    if score >= 126: return AMBER
    if score >= 59:  return colors.HexColor("#F9A825")
    return colors.HexColor("#66BB6A")

def risk_label(score):
    if score >= 277: return "Critical"
    if score >= 199: return "High"
    if score >= 126: return "Moderate"
    if score >= 59:  return "Low"
    return "Stable"

def _table_style(stripe=True):
    base = [
        ("BACKGROUND",    (0,0), (-1,0),  NAVY),
        ("TEXTCOLOR",     (0,0), (-1,0),  WHITE),
        ("FONTNAME",      (0,0), (-1,0),  "Helvetica-Bold"),
        ("FONTSIZE",      (0,0), (-1,0),  8),
        ("TOPPADDING",    (0,0), (-1,0),  7),
        ("BOTTOMPADDING", (0,0), (-1,0),  7),
        ("LEFTPADDING",   (0,0), (-1,-1), 8),
        ("RIGHTPADDING",  (0,0), (-1,-1), 8),
        ("FONTSIZE",      (0,1), (-1,-1), 8.5),
        ("TOPPADDING",    (0,1), (-1,-1), 5),
        ("BOTTOMPADDING", (0,1), (-1,-1), 5),
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("LINEBELOW",     (0,0), (-1,-1), 0.3, GREY_LINE),
        ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHITE, GREY_BG]) if stripe else None,
    ]
    return TableStyle([s for s in base if s is not None])

def _hdr(text):
    return _p(f"<b>{text}</b>", _style("hdr", textColor=WHITE, fontSize=8, fontName="Helvetica-Bold"))

def _col(text, color, bold=True):
    tag = f'<font color="{color.hexval()}">{"<b>" if bold else ""}{text}{"</b>" if bold else ""}</font>'
    return _p(tag, _style("col"))

def _section_header(story, title, subtitle=""):
    story.append(Spacer(1, 5*mm))
    story.append(_p(title, _style("h2", fontSize=15, textColor=NAVY,
                                   fontName="Helvetica-Bold", spaceAfter=2)))
    story.append(HRFlowable(width="100%", thickness=3, color=NAVY,
                             spaceAfter=4 if subtitle else 8))
    if subtitle:
        story.append(_p(subtitle, _style("sub", fontSize=9, textColor=GREY_MID, spaceAfter=8)))

def _stat_card(label, value, sub, accent):
    label_s = _style("sl", fontSize=7,  textColor=GREY_MID, fontName="Helvetica", leading=10)
    value_s = _style("sv", fontSize=20, textColor=accent,   fontName="Helvetica-Bold", leading=24)
    sub_s   = _style("ss", fontSize=7,  textColor=GREY_MID, fontName="Helvetica", leading=10)
    inner = Table([
        [_p(label.upper(), label_s)],
        [_p(str(value), value_s)],
        [_p(sub, sub_s)],
    ], colWidths=[38*mm])
    inner.setStyle(TableStyle([
        ("LEFTPADDING",  (0,0),(-1,-1), 0), ("RIGHTPADDING", (0,0),(-1,-1), 0),
        ("TOPPADDING",   (0,0),(-1,-1), 1), ("BOTTOMPADDING",(0,0),(-1,-1), 1),
    ]))
    outer = Table([[inner]], colWidths=[42*mm])
    outer.setStyle(TableStyle([
        ("BACKGROUND",   (0,0),(-1,-1), GREY_BG),
        ("BOX",          (0,0),(-1,-1), 0.5, GREY_LINE),
        ("LINEBEFORE",   (0,0),(0,-1),  4,   accent),
        ("TOPPADDING",   (0,0),(-1,-1), 10), ("BOTTOMPADDING",(0,0),(-1,-1), 10),
        ("LEFTPADDING",  (0,0),(-1,-1), 10), ("RIGHTPADDING", (0,0),(-1,-1), 6),
    ]))
    return outer

def _footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(GREY_MID)
    canvas.drawString(20*mm, 10*mm,
                      "Malawi Food Security Monitor  —  WFP VAM Data  —  Confidential")
    canvas.drawRightString(190*mm, 10*mm, f"Page {doc.page}")
    canvas.restoreState()

def _build_cover(story, summary, generated_at, analyst_name="WFP VAM System"):
    page_w      = 170*mm
    badge_style = _style("badge", fontSize=9, fontName="Helvetica-Bold",
                          textColor=NAVY, backColor=GOLD, leftIndent=6, rightIndent=6)
    title_style = _style("ctitle", fontSize=28, textColor=WHITE,
                          fontName="Helvetica-Bold", leading=34, spaceAfter=4)
    sub_style   = _style("csub", fontSize=13, textColor=BLUE_LIGHT,
                          fontName="Helvetica", leading=18, spaceAfter=24)
    meta_style  = _style("cmeta", fontSize=10, textColor=BLUE_LIGHT,
                          fontName="Helvetica", leading=18)

    regions = summary.get("total_regions", "3")
    inner = [
        [_p("CONFIDENTIAL ANALYTICAL REPORT", badge_style)],
        [Spacer(1, 20)],
        [_p("Malawi Food Security Monitor", title_style)],
        [_p("Food Price Spike Analysis &amp; District Risk Assessment", sub_style)],
        [HRFlowable(width=60, thickness=4, color=GOLD, spaceAfter=16)],
        [Spacer(1, 8)],
        [_p(f'<font color="#FFFFFF"><b>Analysis Period:</b></font>'
            f'<font color="#A8C4E8">  January 2020 – {datetime.now(tz=timezone.utc).strftime("%B %Y")}</font>',
            meta_style)],
        [_p(f'<font color="#FFFFFF"><b>Data Source:</b></font>'
            f'<font color="#A8C4E8">  World Food Programme (WFP) VAM Food Price Database</font>',
            meta_style)],
        [_p(f'<font color="#FFFFFF"><b>Geographic Scope:</b></font>'
            f'<font color="#A8C4E8">  {summary["total_districts"]} Districts · '
            f'{summary["total_markets"]} Markets · {regions} Regions</font>', meta_style)],
        [_p(f'<font color="#FFFFFF"><b>Report Generated:</b></font>'
            f'<font color="#A8C4E8">  {generated_at} UTC</font>', meta_style)],
        [_p(f'<font color="#FFFFFF"><b>Prepared by:</b></font>'
            f'<font color="#A8C4E8">  {analyst_name}</font>', meta_style)],
        [_p(f'<font color="#FFFFFF"><b>Classification:</b></font>'
            f'<font color="#A8C4E8">  For Official Use — NGO / Government Partners</font>',
            meta_style)],
    ]
    inner_t = Table(inner, colWidths=[140*mm])
    inner_t.setStyle(TableStyle([
        ("LEFTPADDING",  (0,0),(-1,-1), 0), ("RIGHTPADDING", (0,0),(-1,-1), 0),
        ("TOPPADDING",   (0,0),(-1,-1), 2), ("BOTTOMPADDING",(0,0),(-1,-1), 2),
    ]))
    cover_t = Table([[inner_t]], colWidths=[page_w])
    cover_t.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(-1,-1), NAVY),
        ("TOPPADDING",    (0,0),(-1,-1), 50), ("BOTTOMPADDING", (0,0),(-1,-1), 50),
        ("LEFTPADDING",   (0,0),(-1,-1), 30), ("RIGHTPADDING",  (0,0),(-1,-1), 20),
        ("ROWHEIGHT",     (0,0),(-1,-1), 200*mm),
    ]))
    story.append(cover_t)
    story.append(PageBreak())


# ═══════════════════════════════════════════════════════════════════════════════
# Narrative generation helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _severity_sentence(crit, sev, mod, total):
    """Generate a plain-English severity distribution sentence."""
    rate = round((crit + sev + mod) / max(total, 1) * 100, 1)
    if crit == 0:
        return f"No critical spike events were recorded in this period. {mod + sev} moderate-to-severe events were detected ({rate}% spike rate)."
    if crit <= 3:
        return (f"{crit} critical price spike event{'s' if crit != 1 else ''} {'were' if crit != 1 else 'was'} "
                f"detected alongside {sev} severe and {mod} moderate events — "
                f"a combined spike rate of {rate}% of all observations.")
    return (f"{crit} critical price spike events were detected — a severe stress signal "
            f"accompanied by {sev} severe and {mod} moderate events, "
            f"yielding a {rate}% overall spike rate that substantially exceeds the 5% early-warning threshold.")

def _commodity_narrative(commodities, selected_list):
    """Generate commodity-specific narrative for selected commodities."""
    if not commodities:
        return "No commodity-level price data was available for the selected period."
    lines = []
    for c in commodities:
        if selected_list and c["commodity"] not in selected_list and selected_list != ["All commodities"]:
            continue
        pct = float(c.get("avg_pct_change") or 0)
        price = int(float(c.get("avg_price") or 0))
        obs = int(c.get("obs") or 0)
        direction = "rose" if pct > 0 else "fell"
        stress = "at critical stress levels" if pct > 40 else "at elevated levels" if pct > 15 else "within normal range"
        lines.append(
            f"<b>{c['commodity']}</b> averaged MWK {price:,}/KG across {obs} observations, "
            f"with prices {direction} {abs(pct):.1f}% on average — {stress}."
        )
    return " ".join(lines) if lines else "No matching commodity data found for the selected filters."

def _market_narrative(markets):
    """Generate market coverage narrative."""
    if not markets:
        return "No market data was available for this district."
    total_spikes = sum(int(m.get("total_spikes") or 0) for m in markets)
    most_critical = max(markets, key=lambda m: int(m.get("critical_count") or 0))
    crit_count = int(most_critical.get("critical_count") or 0)
    if crit_count > 0:
        return (
            f"{len(markets)} markets are monitored in this district, collectively recording "
            f"{total_spikes} price spike events. {most_critical['market']} experienced the "
            f"highest number of critical events ({crit_count}), suggesting localised supply "
            f"chain disruptions or reduced market competition at that location."
        )
    return (
        f"{len(markets)} markets are monitored in this district, collectively recording "
        f"{total_spikes} spike events. No critical-severity events were detected at any "
        f"individual market — though elevated moderate activity warrants continued monitoring."
    )

def _forecast_narrative(forecast_data):
    """Generate forward-looking narrative from forecast data."""
    if not forecast_data:
        return (
            "No forecast data was available for this report period. "
            "Seasonal patterns suggest continued monitoring is warranted for lean season onset."
        )
    months = forecast_data.get("months", [])
    if not months:
        return "Forecast data was incomplete for the selected period."
    prices = [m.get("predicted_price", 0) for m in months if m.get("predicted_price")]
    if not prices:
        return "Forecast prices could not be computed for this period."
    trend = "rising" if prices[-1] > prices[0] else "declining" if prices[-1] < prices[0] else "stable"
    peak = max(prices)
    peak_month = months[prices.index(peak)].get("month_label", "")
    severity_labels = [m.get("severity", "Normal") for m in months]
    worst = "Normal"
    for s in ["Critical", "Severe", "Moderate"]:
        if s in severity_labels:
            worst = s; break
    return (
        f"The 3-month price forecast indicates a <b>{trend}</b> price trajectory, "
        f"with a projected peak of MWK {int(peak):,}/KG in {peak_month}. "
        f"The worst projected severity level is <b>{worst}</b>. "
        f"{'Immediate preparedness planning is recommended.' if worst in ('Critical','Severe') else 'Continued price monitoring is advised to detect early deterioration.'}"
    )

def _indicator_narrative(indicators, district_name):
    """Generate FSI / PSI indicator narrative."""
    if not indicators:
        return f"Composite indicator data was not available for {district_name}."
    risk_score = float(indicators.get("risk_score") or 0)
    spike_rate = float(indicators.get("spike_rate_pct") or 0)
    rl = risk_label(risk_score)
    # FSI: higher risk_score = lower food security. Express as risk index.
    risk_index = round(risk_score / 436 * 100, 1)
    coverage = indicators.get("monitoring_status", "Unknown")
    return (
        f"{district_name} District has a composite risk score of {int(risk_score)} — "
        f"classified as <b>{rl}</b> ({risk_index:.0f}/100 on the risk index, where 100 represents "
        f"the highest observed national risk). The overall price spike rate stands at "
        f"{spike_rate:.1f}% of observations. "
        f"Market monitoring coverage is rated <b>{coverage}</b>. "
        f"{'This coverage gap means food crises may emerge without adequate early warning — expanding monitoring is a priority.' if 'GAP' in str(coverage).upper() else 'Monitoring coverage is sufficient to detect emerging price pressures.'}"
    )

def _early_warning_narrative(indicators, forecast_data):
    """Generate early warning section narrative."""
    warnings = []
    if indicators:
        spike_rate = float(indicators.get("spike_rate_pct") or 0)
        if spike_rate > 10:
            warnings.append(f"Spike rate of {spike_rate:.1f}% significantly exceeds the 10% high-alert threshold.")
        risk = float(indicators.get("risk_score") or 0)
        if risk > 199:
            warnings.append(f"Composite risk score of {int(risk)} places this district in the top-priority intervention category.")
    if forecast_data:
        months = forecast_data.get("months", [])
        for m in months:
            if m.get("severity") in ("Critical", "Severe"):
                warnings.append(
                    f"Forecast projects {m['severity'].lower()}-severity prices for "
                    f"{m.get('month_label','the coming period')} — immediate preparedness action required."
                )
                break
    if not warnings:
        return "No acute early warning signals were identified for this reporting period. Routine monitoring should continue."
    return " ".join(warnings) + " Response planning should begin immediately."


# ═══════════════════════════════════════════════════════════════════════════════
# PDF builder — District
# ═══════════════════════════════════════════════════════════════════════════════

def build_district_pdf(data: dict) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=16*mm,  bottomMargin=18*mm,
        title=f"District Report — {data['district']['district']}",
        author="WFP VAM",
    )

    story        = []
    district     = data["district"]
    spikes       = data["spikes"]
    prices       = data["prices"]
    markets      = data["markets"]
    indicators   = data.get("indicators")
    forecast     = data.get("forecast")
    date_range   = data.get("date_range", {})
    commodities  = data.get("commodity_list", ["All commodities"])
    commodity_str = ", ".join(commodities) if commodities else "All commodities"
    generated_at = datetime.now(tz=timezone.utc).strftime("%d %B %Y %H:%M")

    body    = _style("body",  fontSize=9,   textColor=GREY_DARK, leading=14)
    bodybig = _style("body2", fontSize=9.5, textColor=GREY_DARK, leading=15)
    small   = _style("small", fontSize=8,   textColor=GREY_MID,  leading=11)
    lead    = _style("lead",  fontSize=10.5, textColor=GREY_DARK, leading=16,
                     fontName="Helvetica", spaceAfter=4)

    rc = risk_color(district["risk_score"])
    rl = risk_label(district["risk_score"])
    crit  = int(district.get("critical_count") or 0)
    sev   = int(district.get("severe_count") or 0)
    mod   = int(district.get("moderate_count") or 0)
    total_spikes_district = crit * 3 + sev * 2 + mod  # rough obs proxy

    # ── Cover ──────────────────────────────────────────────────────────────────
    _build_cover(story, {
        "total_districts": 1,
        "total_markets"  : len(markets),
        "total_regions"  : district.get("region", "Southern"),
        "highest_risk_district": {
            "name"       : district["district"],
            "region"     : district["region"],
            "risk_score" : district["risk_score"],
        },
        "most_spiked_commodity": {"name": commodity_str, "critical_events": crit},
        "monitoring_gaps"      : {"critical_gap_districts": 0},
        "spike_events"         : {"total": total_spikes_district, "critical": crit,
                                   "severe": sev, "moderate": mod},
    }, generated_at)

    # ── Section 1: District Overview ───────────────────────────────────────────
    _section_header(
        story, f"1. District Overview — {district['district']}",
        f"{district['region']} Region  ·  "
        f"{date_range.get('from','2020-01-01')} to {date_range.get('to','today')}  ·  "
        f"Commodity filter: {commodity_str}"
    )

    # stat cards
    cards = Table([[
        _stat_card("Risk Score",      int(district["risk_score"]),     rl, rc),
        _stat_card("Critical Spikes", crit,                             "spike events", RED),
        _stat_card("Avg Price",
                   f"{int(float(district['avg_price'])):,} MWK",
                   "across all commodities", NAVY),
        _stat_card("Markets", len(markets), "monitored markets", GREEN),
    ]], colWidths=[42*mm]*4)
    cards.setStyle(TableStyle([
        ("LEFTPADDING",  (0,0),(-1,-1), 0), ("RIGHTPADDING", (0,0),(-1,-1), 0),
        ("TOPPADDING",   (0,0),(-1,-1), 0), ("BOTTOMPADDING",(0,0),(-1,-1), 0),
        ("INNERGRID",    (0,0),(-1,-1), 2, WHITE),
    ]))
    story.append(cards)
    story.append(Spacer(1, 5*mm))

    # situation narrative
    situation_text = (
        f"{district['district']} District, located in Malawi's {district['region']} Region, "
        f"presents a food security risk profile classified as <b>{rl}</b> with a composite "
        f"risk score of <b>{int(district['risk_score'])}</b>. "
        f"{_severity_sentence(crit, sev, mod, total_spikes_district)} "
        f"The district's average market price of MWK {int(float(district['avg_price'])):,}/KG "
        f"reflects the cumulative impact of price pressures over the analysis period."
    )
    story.append(NarrativeBox(
        "Situation Summary",
        [_p(situation_text, bodybig)],
        170*mm, bg=BLUE_BG, heading_bg=NAVY
    ))
    story.append(Spacer(1, 5*mm))

    # ── Section 2: Food Security Indicators ───────────────────────────────────
    if indicators:
        _section_header(story, "2. Food Security Indicators",
                        "Composite risk score, FSI, price stress index, and monitoring coverage")

        risk_index = round(float(indicators.get("risk_score") or 0) / 436 * 100, 1)
        spike_rate = float(indicators.get("spike_rate_pct") or 0)
        coverage   = str(indicators.get("monitoring_status") or "Unknown")

        ind_rows = [
            [_hdr("Indicator"), _hdr("Value"), _hdr("Threshold"), _hdr("Status")],
            [_p("Composite Risk Score", body), _col(str(int(district["risk_score"])), rc),
             _p("< 59 = Stable", small), _col(rl, rc)],
            [_p("Risk Index (0–100)", body),
             _col(f"{risk_index:.0f}/100", rc),
             _p("100 = highest national risk", small), _col(rl, rc)],
            [_p("Price Spike Rate", body),
             _col(f"{spike_rate:.1f}%",
                  RED if spike_rate > 10 else AMBER if spike_rate > 5 else GREEN),
             _p("> 10% = High Alert", small),
             _col("HIGH ALERT" if spike_rate > 10 else "ELEVATED" if spike_rate > 5 else "NORMAL",
                  RED if spike_rate > 10 else AMBER if spike_rate > 5 else GREEN)],
            [_p("Market Monitoring Coverage", body),
             _p(coverage, body),
             _p("Adequate = sufficient early warning", small),
             _col("GAP" if "GAP" in coverage.upper() else "OK",
                  RED if "GAP" in coverage.upper() else GREEN)],
        ]
        ind_t = Table(ind_rows, colWidths=[55*mm, 35*mm, 55*mm, 25*mm], repeatRows=1)
        ind_t.setStyle(_table_style())
        story.append(ind_t)
        story.append(Spacer(1, 4*mm))

        # indicator narrative
        story.append(LeftBorderBox(
            _p(_indicator_narrative(indicators, district["district"]),
               _style("ind", fontSize=9.5, textColor=GREY_DARK, leading=14)),
            170*mm, BLUE_BG, NAVY, border_width=4
        ))

        # early warning
        ew_text = _early_warning_narrative(indicators, forecast)
        ew_color = RED if ("immediately" in ew_text.lower() or "immediately" in ew_text.lower()) else AMBER
        story.append(Spacer(1, 4*mm))
        story.append(LeftBorderBox(
            _p(f"<b>Early Warning Signal:</b> {ew_text}",
               _style("ew", fontSize=9.5, textColor=GREY_DARK, leading=14)),
            170*mm, ORANGE_BG, ew_color, border_width=4
        ))
        story.append(PageBreak())

    # ── Section 3: Commodity Price Analysis ───────────────────────────────────
    if prices:
        _section_header(story, "3. Commodity Price Analysis",
                        "Average prices, percentage change, and price stress classification")

        # narrative
        comm_narr = _commodity_narrative(
            [dict(p) for p in prices], commodities
        )
        story.append(LeftBorderBox(
            _p(comm_narr, _style("cn", fontSize=9.5, textColor=GREY_DARK, leading=14)),
            170*mm, GREEN_BG, GREEN, border_width=4
        ))
        story.append(Spacer(1, 4*mm))

        price_rows = [[_hdr("Commodity"), _hdr("Avg Price (MWK/KG)"),
                       _hdr("Avg % Change"), _hdr("Stress Level"), _hdr("Observations")]]
        for p in prices:
            pct = float(p.get("avg_pct_change") or 0)
            stress = "Critical" if pct > 40 else "Elevated" if pct > 15 else "Normal"
            sc = RED if pct > 40 else AMBER if pct > 15 else GREEN
            price_rows.append([
                _p(f"<b>{p['commodity']}</b>", body),
                _p(f"{int(float(p['avg_price'])):,} MWK", body),
                _col(f"{'+' if pct >= 0 else ''}{pct:.1f}%", sc),
                _col(stress, sc),
                _p(str(p["obs"]), small),
            ])
        price_t = Table(price_rows, colWidths=[55*mm, 35*mm, 30*mm, 30*mm, 20*mm], repeatRows=1)
        price_t.setStyle(_table_style())
        story.append(price_t)
        story.append(Spacer(1, 5*mm))

    # ── Section 4: 3-Month Price Forecast ─────────────────────────────────────
    _section_header(story, "4. Price Forecast Outlook",
                    "3-month forward projection — MFPI model · confidence band: 80%")

    forecast_narr = _forecast_narrative(forecast)
    story.append(LeftBorderBox(
        _p(forecast_narr, _style("fn2", fontSize=9.5, textColor=GREY_DARK, leading=14)),
        170*mm, ORANGE_BG, ORANGE, border_width=4
    ))
    story.append(Spacer(1, 4*mm))

    if forecast and forecast.get("months"):
        fc_rows = [[_hdr("Month"), _hdr("Season"), _hdr("Projected Price (MWK/KG)"),
                    _hdr("Change vs Baseline"), _hdr("Severity")]]
        for m in forecast["months"]:
            sc = {"Critical": RED, "Severe": ORANGE, "Moderate": AMBER,
                  "Normal": GREEN}.get(m.get("severity", "Normal"), GREEN)
            fc_rows.append([
                _p(f"<b>{m.get('month_label','—')}</b>", body),
                _p(m.get("season", "—"), small),
                _col(f"{int(m.get('predicted_price',0)):,} MWK", sc),
                _col(f"{'+' if float(m.get('pct_vs_baseline',0)) >= 0 else ''}"
                     f"{float(m.get('pct_vs_baseline',0)):.1f}%", sc),
                _col(m.get("severity", "Normal"), sc),
            ])
        fc_t = Table(fc_rows, colWidths=[35*mm, 30*mm, 45*mm, 35*mm, 25*mm], repeatRows=1)
        fc_t.setStyle(_table_style())
        story.append(fc_t)
    else:
        story.append(_p(
            "Detailed month-by-month forecast data was not available for this district and period. "
            "The MFPI model requires a minimum of 12 months of price history per commodity.",
            _style("nofc", fontSize=9, textColor=GREY_MID, leading=13)
        ))
    story.append(PageBreak())

    # ── Section 5: Market Coverage ─────────────────────────────────────────────
    if markets:
        _section_header(story, "5. Market Coverage",
                        "Monitored markets — spike frequency and monitoring health")

        mkt_narr = _market_narrative([dict(m) for m in markets])
        story.append(LeftBorderBox(
            _p(mkt_narr, _style("mn", fontSize=9.5, textColor=GREY_DARK, leading=14)),
            170*mm, GREY_BG, NAVY, border_width=4
        ))
        story.append(Spacer(1, 4*mm))

        mkt_rows = [[_hdr("Market"), _hdr("Commodities"), _hdr("Total Spikes"),
                     _hdr("Critical"), _hdr("Spike Rate"), _hdr("Last Report")]]
        for m in markets:
            # FIX: strip timestamp to date only
            last = str(m.get("latest_date") or "—")
            if "T" in last:
                last = last.split("T")[0]
            elif " " in last and len(last) > 10:
                last = last[:10]
            mkt_rows.append([
                _p(f"<b>{m['market']}</b>", body),
                _p(str(m["num_commodities"]), small),
                _p(str(m["total_spikes"]), body),
                _col(str(m["critical_count"]), RED if int(m["critical_count"]) > 0 else GREY_MID),
                _p(f"{float(m['spike_rate_pct']):.1f}%", small),
                _p(last, small),
            ])
        mkt_t = Table(mkt_rows,
                      colWidths=[45*mm, 22*mm, 25*mm, 20*mm, 22*mm, 36*mm],
                      repeatRows=1)
        mkt_t.setStyle(_table_style())
        story.append(mkt_t)
        story.append(PageBreak())

    # ── Section 6: Spike Events ────────────────────────────────────────────────
    if spikes:
        _section_header(
            story, "6. Price Spike Events",
            f"Detected spike events · {date_range.get('from','2020-01-01')} to {date_range.get('to','today')}"
        )

        spike_intro = (
            f"The following table records all detected price spike events for {district['district']} "
            f"District during the selected period, ordered by date (most recent first). "
            f"Events are classified using the WFP ALPS dual-method framework requiring both "
            f"month-over-month price change ≥20% AND z-score ≥1.5 above the 12-month rolling mean."
        )
        story.append(LeftBorderBox(
            _p(spike_intro, _style("si", fontSize=9, textColor=GREY_DARK, leading=13)),
            170*mm, GREY_BG, NAVY, border_width=3
        ))
        story.append(Spacer(1, 4*mm))

        sev_colors  = {"Critical": RED, "Severe": ORANGE, "Moderate": AMBER}
        spike_rows  = [[_hdr("Date"), _hdr("Market"), _hdr("Commodity"),
                        _hdr("Price (MWK/KG)"), _hdr("% Jump"), _hdr("Z-Score"),
                        _hdr("Severity")]]
        for s in spikes:
            sc = sev_colors.get(s["spike_severity"], GREEN)
            spike_rows.append([
                _p(str(s["date"]), small),
                _p(s["market"], body),
                _p(f"<b>{s['commodity']}</b>", body),
                _p(f"{int(float(s['price'])):,} MWK", body),
                _col(f"+{float(s['pct_change']):.1f}%", sc),
                _p(f"{float(s.get('zscore',0)):.2f}", small),
                _col(s["spike_severity"], sc),
            ])
        spike_t = Table(spike_rows,
                        colWidths=[22*mm,35*mm,28*mm,25*mm,18*mm,18*mm,24*mm],
                        repeatRows=1)
        spike_t.setStyle(_table_style())
        story.append(spike_t)

    # ── Section 7: Recommendations ─────────────────────────────────────────────
    story.append(Spacer(1, 6*mm))
    _section_header(story, "7. Recommendations")

    risk_score = district["risk_score"]
    if risk_score >= 277:
        rec_text = (
            f"<b>IMMEDIATE ACTION REQUIRED.</b> {district['district']} District's Critical risk "
            f"classification demands immediate food assistance intervention. Activate emergency "
            f"food pipelines and commission a rapid market assessment. Establish maize price "
            f"monitoring thresholds at 1,200 MWK/KG to trigger automatic alert escalation to "
            f"DoDMA and WFP country offices. Expand market monitoring coverage given the current gap status."
        )
        rec_color = RED; rec_bg = ORANGE_BG
    elif risk_score >= 126:
        rec_text = (
            f"<b>ELEVATED MONITORING REQUIRED.</b> {district['district']} District requires "
            f"enhanced monitoring frequency and pre-positioning of contingency food stocks. "
            f"Assess supply chain constraints contributing to elevated price volatility. "
            f"Review monitoring coverage gaps to ensure early warning capability."
        )
        rec_color = AMBER; rec_bg = ORANGE_BG
    else:
        rec_text = (
            f"<b>ROUTINE MONITORING.</b> {district['district']} District is currently operating "
            f"within acceptable food security parameters. Maintain standard monitoring protocols "
            f"and continue seasonal baseline assessments. Flag any sudden deterioration in market "
            f"integration or coverage for review."
        )
        rec_color = GREEN; rec_bg = GREEN_BG

    story.append(LeftBorderBox(
        _p(rec_text, _style("rec", fontSize=10, textColor=GREY_DARK, leading=15)),
        170*mm, rec_bg, rec_color, border_width=5
    ))

    # ── Footer box ─────────────────────────────────────────────────────────────
    story.append(Spacer(1, 8*mm))
    story.append(LeftBorderBox(
        _p(
            f"<b>District:</b> {district['district']}  |  "
            f"<b>Generated:</b> {generated_at} UTC  |  "
            f"<b>Source:</b> WFP VAM Food Price Database  |  "
            f"<b>Methodology:</b> Dual-method ALPS (% change + z-score)",
            _style("fn_box", fontSize=8, textColor=GREY_MID, leading=12)
        ),
        170*mm, BLUE_BG, NAVY, border_width=3
    ))

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX builder — District (editable Word document)
# ═══════════════════════════════════════════════════════════════════════════════

def _docx_heading(doc, text, level=1, color_hex="1B3A6B"):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return h

def _docx_para(doc, text, bold=False, size=10, color_hex=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    p.paragraph_format.space_after = Pt(6)
    return p

def _docx_callout(doc, label, text, color_hex="1B3A6B"):
    """Add a left-bordered callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color_hex)
    run.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    # left border via pPr
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def _docx_table(doc, headers, rows, col_widths_inches):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1B3A6B')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
            if ri % 2 == 0:
                tc = cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F8F9FA')
                tcPr.append(shd)
    # column widths
    for i, w in enumerate(col_widths_inches):
        for row in table.rows:
            row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def build_district_docx(data: dict) -> bytes:
    """Build an editable Word document with the full district situation report."""
    district     = data["district"]
    spikes       = data["spikes"]
    prices       = data["prices"]
    markets      = data["markets"]
    indicators   = data.get("indicators")
    forecast     = data.get("forecast")
    date_range   = data.get("date_range", {})
    commodities  = data.get("commodity_list", ["All commodities"])
    commodity_str = ", ".join(commodities) if commodities else "All commodities"
    generated_at = datetime.now(tz=timezone.utc).strftime("%d %B %Y %H:%M")

    rc_hex = {
        "Critical": "B71C1C", "High": "EF5350",
        "Moderate": "F9A825", "Low":  "F9A825", "Stable": "2E7D32"
    }
    rl = risk_label(district["risk_score"])
    crit = int(district.get("critical_count") or 0)
    sev  = int(district.get("severe_count") or 0)
    mod  = int(district.get("moderate_count") or 0)
    total_obs = crit * 3 + sev * 2 + mod

    doc = DocxDocument()

    # page setup: A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── Title block ─────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title_p.add_run("MALAWI FOOD SECURITY MONITOR")
    tr.bold = True; tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(27, 58, 107)

    subtitle_p = doc.add_paragraph()
    sr = subtitle_p.add_run(f"District Situation Report — {district['district']}")
    sr.font.size = Pt(14); sr.font.color.rgb = RGBColor(102, 102, 102)

    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Period: {date_range.get('from','2020-01-01')} to {date_range.get('to','today')}  ·  "
                   f"Commodity: {commodity_str}  ·  Generated: {generated_at} UTC").font.size = Pt(9)
    meta_p.paragraph_format.space_after = Pt(12)

    # horizontal rule simulation
    border_p = doc.add_paragraph()
    pPr = border_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '1B3A6B')
    pBdr.append(bottom)
    pPr.append(pBdr)
    border_p.paragraph_format.space_after = Pt(12)

    # ── 1. Situation Summary ────────────────────────────────────────────────
    _docx_heading(doc, "1. Situation Summary")
    situation_text = (
        f"{district['district']} District, located in Malawi's {district['region']} Region, "
        f"presents a food security risk profile classified as {rl} with a composite risk score "
        f"of {int(district['risk_score'])}. "
        f"{_severity_sentence(crit, sev, mod, total_obs)} "
        f"The district's average market price of MWK {int(float(district['avg_price'])):,}/KG "
        f"reflects the cumulative impact of price pressures over the analysis period."
    )
    _docx_callout(doc, "Assessment", situation_text, rc_hex.get(rl, "1B3A6B"))

    # key stats table
    _docx_heading(doc, "Key Statistics", level=2)
    _docx_table(doc,
        ["Indicator", "Value", "Classification"],
        [
            ["Composite Risk Score", str(int(district["risk_score"])), rl],
            ["Critical Spike Events", str(crit), "Immediate action" if crit > 5 else "Monitor"],
            ["Average Price (MWK/KG)", f"{int(float(district['avg_price'])):,}", "—"],
            ["Markets Monitored", str(len(markets)), "—"],
        ],
        [2.5, 2.0, 2.5]
    )

    # ── 2. Food Security Indicators ─────────────────────────────────────────
    if indicators:
        _docx_heading(doc, "2. Food Security Indicators")
        _docx_para(doc, _indicator_narrative(indicators, district["district"]))

        risk_index = round(float(indicators.get("risk_score") or 0) / 436 * 100, 1)
        spike_rate = float(indicators.get("spike_rate_pct") or 0)
        coverage   = str(indicators.get("monitoring_status") or "Unknown")

        _docx_table(doc,
            ["Indicator", "Value", "Threshold", "Status"],
            [
                ["Risk Index (0–100)", f"{risk_index:.0f}/100", "100 = worst nationally", rl],
                ["Price Spike Rate", f"{spike_rate:.1f}%", "> 10% = High Alert",
                 "HIGH ALERT" if spike_rate > 10 else "ELEVATED" if spike_rate > 5 else "NORMAL"],
                ["Monitoring Coverage", coverage, "Adequate = sufficient warning",
                 "⚠ GAP" if "GAP" in coverage.upper() else "OK"],
            ],
            [2.5, 1.8, 2.8, 1.6]
        )

        ew_text = _early_warning_narrative(indicators, forecast)
        _docx_callout(doc, "Early Warning", ew_text, "E65100")

    # ── 3. Commodity Analysis ───────────────────────────────────────────────
    if prices:
        _docx_heading(doc, "3. Commodity Price Analysis")
        comm_narr = _commodity_narrative([dict(p) for p in prices], commodities)
        # strip HTML tags for docx
        import re
        clean_narr = re.sub(r'<[^>]+>', '', comm_narr)
        _docx_para(doc, clean_narr)

        price_rows = []
        for p in prices:
            pct = float(p.get("avg_pct_change") or 0)
            stress = "Critical" if pct > 40 else "Elevated" if pct > 15 else "Normal"
            price_rows.append([
                p["commodity"],
                f"{int(float(p['avg_price'])):,} MWK",
                f"{'+' if pct >= 0 else ''}{pct:.1f}%",
                stress,
                str(p["obs"]),
            ])
        _docx_table(doc,
            ["Commodity", "Avg Price (MWK/KG)", "Avg % Change", "Stress Level", "Observations"],
            price_rows,
            [2.2, 1.8, 1.5, 1.5, 1.2]
        )

    # ── 4. Price Forecast ───────────────────────────────────────────────────
    _docx_heading(doc, "4. 3-Month Price Forecast")
    import re
    clean_fc_narr = re.sub(r'<[^>]+>', '', _forecast_narrative(forecast))
    _docx_callout(doc, "Forecast Summary", clean_fc_narr, "E65100")

    if forecast and forecast.get("months"):
        fc_rows = []
        for m in forecast["months"]:
            fc_rows.append([
                m.get("month_label", "—"),
                m.get("season", "—"),
                f"{int(m.get('predicted_price', 0)):,} MWK",
                f"{'+' if float(m.get('pct_vs_baseline', 0)) >= 0 else ''}"
                f"{float(m.get('pct_vs_baseline', 0)):.1f}%",
                m.get("severity", "Normal"),
            ])
        _docx_table(doc,
            ["Month", "Season", "Projected Price", "vs Baseline", "Severity"],
            fc_rows,
            [1.5, 1.5, 1.8, 1.5, 1.5]
        )
    else:
        _docx_para(doc,
            "Detailed forecast data was not available. The MFPI model requires a minimum "
            "of 12 months of price history per commodity.", italic=True, color_hex="666666")

    # ── 5. Market Coverage ──────────────────────────────────────────────────
    if markets:
        _docx_heading(doc, "5. Market Coverage")
        import re
        clean_mkt = re.sub(r'<[^>]+>', '', _market_narrative([dict(m) for m in markets]))
        _docx_para(doc, clean_mkt)

        mkt_rows = []
        for m in markets:
            last = str(m.get("latest_date") or "—")
            if "T" in last: last = last.split("T")[0]
            elif " " in last and len(last) > 10: last = last[:10]
            mkt_rows.append([
                m["market"],
                str(m["num_commodities"]),
                str(m["total_spikes"]),
                str(m["critical_count"]),
                f"{float(m['spike_rate_pct']):.1f}%",
                last,
            ])
        _docx_table(doc,
            ["Market", "Commodities", "Total Spikes", "Critical", "Spike Rate", "Last Report"],
            mkt_rows,
            [2.0, 1.2, 1.3, 1.0, 1.2, 1.5]
        )

    # ── 6. Spike Events ─────────────────────────────────────────────────────
    if spikes:
        _docx_heading(doc, "6. Price Spike Events")
        _docx_para(doc,
            f"The following table records all detected price spike events for "
            f"{district['district']} District during the selected period, classified using "
            f"the WFP ALPS dual-method framework (% change + z-score).")

        spike_rows = []
        for s in spikes:
            spike_rows.append([
                str(s["date"]),
                s["market"],
                s["commodity"],
                f"{int(float(s['price'])):,}",
                f"+{float(s['pct_change']):.1f}%",
                f"{float(s.get('zscore', 0)):.2f}",
                s["spike_severity"],
            ])
        _docx_table(doc,
            ["Date", "Market", "Commodity", "Price (MWK)", "% Jump", "Z-Score", "Severity"],
            spike_rows,
            [1.2, 1.8, 1.5, 1.2, 1.0, 1.0, 1.2]
        )

    # ── 7. Recommendations ─────────────────────────────────────────────────
    _docx_heading(doc, "7. Recommendations")
    risk_score = district["risk_score"]
    if risk_score >= 277:
        rec = (
            f"IMMEDIATE ACTION REQUIRED. {district['district']} District's Critical risk "
            "classification demands immediate food assistance intervention. Activate emergency "
            "food pipelines and commission a rapid market assessment. Establish maize price "
            "monitoring thresholds at 1,200 MWK/KG to trigger automatic alert escalation."
        )
        rec_color = "B71C1C"
    elif risk_score >= 126:
        rec = (
            f"ELEVATED MONITORING REQUIRED. {district['district']} District requires "
            "enhanced monitoring frequency and pre-positioning of contingency food stocks. "
            "Assess supply chain constraints contributing to elevated price volatility."
        )
        rec_color = "E65100"
    else:
        rec = (
            f"ROUTINE MONITORING. {district['district']} District is operating within "
            "acceptable food security parameters. Maintain standard monitoring protocols "
            "and continue seasonal baseline assessments."
        )
        rec_color = "2E7D32"
    _docx_callout(doc, "Recommended Action", rec, rec_color)

    # ── Footer ──────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(24)
    fr = footer_p.add_run(
        f"Report generated by Malawi Food Security Monitor  |  "
        f"Source: WFP VAM Food Price Database  |  "
        f"Generated: {generated_at} UTC  |  "
        f"Methodology: WFP ALPS dual-method spike detection"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(102, 102, 102)
    fr.italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# National PDF builder (unchanged structure, adds narrative callouts)
# ═══════════════════════════════════════════════════════════════════════════════

def build_pdf(data: dict) -> bytes:
    """National report — all 28 districts."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=16*mm,  bottomMargin=18*mm,
        title="Malawi Food Security Monitor", author="WFP VAM",
    )

    story        = []
    summary      = data["summary"]
    districts    = data["districts"]
    crit_spikes  = data["critical_spikes"]
    commodities  = data["commodities"]
    gaps         = data["gaps"]
    generated_at = datetime.now(tz=timezone.utc).strftime("%d %B %Y %H:%M")

    total = summary["spike_events"]["total"]    or 1
    crit  = summary["spike_events"]["critical"] or 0
    sev   = summary["spike_events"]["severe"]   or 0
    mod   = summary["spike_events"]["moderate"] or 0

    body  = _style("body",  fontSize=9, textColor=GREY_DARK, leading=13)
    small = _style("small", fontSize=8, textColor=GREY_MID,  leading=11)

    _build_cover(story, summary, generated_at)

    # Executive Summary
    _section_header(story, "Executive Summary",
                    "Key findings from the automated food price spike detection system")

    cards = Table([[
        _stat_card("Markets Monitored", summary["total_markets"],
                   f"{summary['total_districts']} districts covered", NAVY),
        _stat_card("Critical Spikes", crit, f"of {total:,} total events", RED),
        _stat_card("Highest Risk District",
                   summary["highest_risk_district"]["name"],
                   f"Score: {int(summary['highest_risk_district']['risk_score'])}", ORANGE),
        _stat_card("Monitoring Gaps",
                   summary["monitoring_gaps"]["critical_gap_districts"],
                   "critical gap districts", RED),
    ]], colWidths=[42*mm]*4)
    cards.setStyle(TableStyle([
        ("LEFTPADDING",  (0,0),(-1,-1), 0), ("RIGHTPADDING", (0,0),(-1,-1), 0),
        ("TOPPADDING",   (0,0),(-1,-1), 0), ("BOTTOMPADDING",(0,0),(-1,-1), 0),
        ("INNERGRID",    (0,0),(-1,-1), 2, WHITE),
    ]))
    story.append(cards)
    story.append(Spacer(1, 6*mm))

    # Executive narrative
    exec_narr = (
        f"This report covers {summary['total_markets']} monitored markets across "
        f"{summary['total_districts']} districts of Malawi for the period January 2020 – present. "
        f"{_severity_sentence(crit, sev, mod, total)} "
        f"{summary['highest_risk_district']['name']} District in "
        f"{summary['highest_risk_district']['region']} Region holds the highest composite "
        f"risk score ({int(summary['highest_risk_district']['risk_score'])}), driven primarily "
        f"by {summary['most_spiked_commodity']['name']} price volatility "
        f"({summary['most_spiked_commodity']['critical_events']} critical events). "
        f"{summary['monitoring_gaps']['critical_gap_districts']} districts have critical "
        f"monitoring gaps — meaning food crises in those areas may go undetected."
    )
    story.append(NarrativeBox(
        "National Situation Overview",
        [_p(exec_narr, _style("en", fontSize=9.5, textColor=GREY_DARK, leading=14))],
        170*mm, bg=BLUE_BG, heading_bg=NAVY
    ))
    story.append(Spacer(1, 5*mm))

    story.append(LeftBorderBox(
        _p(
            "<b>Methodology:</b> Spike detection uses a dual-method approach requiring both "
            "month-over-month percentage change ≥20% AND rolling 12-month z-score ≥1.5. "
            "Severity follows the WFP ALPS framework (Moderate ≥20%/z1.5 · Severe ≥40%/z2.0 · "
            "Critical ≥60%/z2.5).",
            _style("method", fontSize=9, textColor=GREY_DARK, leading=13)
        ),
        170*mm, BLUE_BG, NAVY, border_width=3
    ))
    story.append(Spacer(1, 4*mm))

    sev_data = [
        [_hdr("Severity Level"), _hdr("Thresholds"), _hdr("Count"),
         _hdr("% of Observations"), _hdr("Recommended Action")],
        [_col("Critical", RED),  "≥60% change · z-score ≥2.5",
         _col(str(crit), RED),   f"{round(crit/total*100,2)}%", "Immediate intervention"],
        [_col("Severe",   ORANGE),"≥40% change · z-score ≥2.0",
         _col(str(sev), ORANGE), f"{round(sev/total*100,2)}%",  "Emergency response"],
        [_col("Moderate", AMBER), "≥20% change · z-score ≥1.5",
         _col(str(mod), AMBER),  f"{round(mod/total*100,2)}%",  "Monitor closely"],
        [_col("Normal",   GREEN), "Below thresholds",
         _col(str(max(0, total-crit-sev-mod)), GREEN), "—", "No action required"],
    ]
    sev_t = Table(sev_data, colWidths=[30*mm, 48*mm, 20*mm, 32*mm, 40*mm], repeatRows=1)
    sev_t.setStyle(_table_style())
    story.append(sev_t)
    story.append(PageBreak())

    # District rankings
    _section_header(story, "District Risk Rankings",
                    "All 28 districts ranked by weighted composite score (Critical×3 + Severe×2 + Moderate×1)")
    dist_rows = [[_hdr("#"), _hdr("District"), _hdr("Region"), _hdr("Risk Level"),
                  _hdr("Risk Score"), _hdr("Critical"), _hdr("Severe"),
                  _hdr("Moderate"), _hdr("Spike Rate")]]
    for i, d in enumerate(districts, 1):
        rc = risk_color(d["risk_score"]); rl = risk_label(d["risk_score"])
        dist_rows.append([
            _p(str(i), small), _p(f"<b>{d['name_1']}</b>", body),
            _p(d["region"], small), _col(rl, rc),
            _col(str(int(d["risk_score"])), rc),
            _col(str(d["critical_count"]), RED),
            _p(str(d["severe_count"]), body),
            _p(str(d["moderate_count"]), body),
            _p(f"{round(float(d['spike_rate_pct']),1)}%", small),
        ])
    dist_t = Table(dist_rows,
                   colWidths=[8*mm,28*mm,24*mm,20*mm,18*mm,17*mm,15*mm,18*mm,18*mm],
                   repeatRows=1)
    dist_t.setStyle(_table_style())
    story.append(dist_t)
    story.append(PageBreak())

    # Critical spikes
    _section_header(story, "Critical Spike Events",
                    "Most recent critical price spike events (price jump ≥60% + z-score ≥2.5)")
    spike_rows = [[_hdr("Date"), _hdr("District"), _hdr("Market"),
                   _hdr("Commodity"), _hdr("Price (MWK/KG)"), _hdr("% Jump")]]
    for s in crit_spikes[:20]:
        spike_rows.append([
            _p(s["date"], small), _p(f"<b>{s['district']}</b>", body),
            _p(s["market"], body), _p(f"<b>{s['commodity']}</b>", body),
            _p(f"{int(float(s['price'])):,} MWK", body),
            _col(f"+{round(float(s['pct_change']),1)}%", RED),
        ])
    spike_t = Table(spike_rows,
                    colWidths=[22*mm,30*mm,38*mm,32*mm,28*mm,20*mm], repeatRows=1)
    spike_t.setStyle(_table_style())
    story.append(spike_t)
    story.append(PageBreak())

    # Commodity analysis
    _section_header(story, "Commodity Analysis",
                    "Commodities with the highest price shock frequency")
    comm_rows = [[_hdr("Commodity"), _hdr("Critical Events"),
                  _hdr("Total Spikes"), _hdr("Spike Rate"), _hdr("Assessment")]]
    for c in commodities:
        crit_c = int(c["critical_count"]); rate = float(c["spike_rate"])
        assessment = ("Priority — immediate attention" if crit_c >= 20 else
                      "Elevated — monitor" if crit_c >= 5 else "Routine")
        comm_rows.append([
            _p(f"<b>{c['commodity']}</b>", body),
            _col(str(crit_c), RED),
            _p(str(c["total_spikes"]), body),
            _p(f"{round(rate,1)}%", body),
            _p(assessment, _style("as", fontSize=8.5, textColor=GREY_DARK, leading=12)),
        ])
    comm_t = Table(comm_rows, colWidths=[55*mm, 30*mm, 28*mm, 22*mm, 35*mm], repeatRows=1)
    comm_t.setStyle(_table_style())
    story.append(comm_t)
    story.append(PageBreak())

    # Monitoring gaps
    _section_header(story, "Market Monitoring Gaps",
                    "Districts with high risk scores but insufficient monitoring coverage")
    story.append(LeftBorderBox(
        _p(
            "<b>Why this matters:</b> A high-risk district with few monitored markets means "
            "food crises may go undetected until they escalate. These gaps represent priority "
            "areas for WFP monitoring expansion and humanitarian pre-positioning.",
            _style("why", fontSize=9, textColor=GREY_DARK, leading=13)
        ),
        170*mm, ORANGE_BG, ORANGE, border_width=4
    ))
    story.append(Spacer(1, 5*mm))

    sc_map = {
        "CRITICAL GAP": RED, "HIGH GAP": ORANGE, "MODERATE GAP": AMBER,
        "ADEQUATE": GREEN,   "NO MONITORING": GREY_MID,
    }
    gap_rows = [[_hdr("District"), _hdr("Region"), _hdr("Risk Score"),
                 _hdr("Markets"), _hdr("Monitoring Status")]]
    for g in gaps:
        sc = sc_map.get(g["monitoring_status"], GREY_MID)
        gap_rows.append([
            _p(f"<b>{g['district']}</b>", body), _p(g["region"], body),
            _col(str(int(g["risk_score"])), RED),
            _p(str(int(g["market_count"])), body),
            _col(g["monitoring_status"], sc),
        ])
    gap_t = Table(gap_rows, colWidths=[38*mm, 35*mm, 28*mm, 22*mm, 47*mm], repeatRows=1)
    gap_t.setStyle(_table_style())
    story.append(gap_t)
    story.append(Spacer(1, 8*mm))

    story.append(LeftBorderBox(
        _p(
            f"<b>Report generated by:</b> Malawi Food Security GIS Monitor  |  "
            f"<b>Data source:</b> WFP VAM Food Price Database  |  "
            f"<b>Generated:</b> {generated_at} UTC",
            _style("fn", fontSize=8, textColor=GREY_MID, leading=12)
        ),
        170*mm, BLUE_BG, NAVY, border_width=3
    ))

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# API Endpoints
# ═══════════════════════════════════════════════════════════════════════════════

@router.get("/generate")
async def generate_report():
    """National report — all 28 districts, PDF only."""
    try:
        pool = await get_pool()
        async with pool.acquire() as conn:
            summary = await conn.fetchrow("""
                SELECT
                    (SELECT COUNT(*) FROM markets)                    AS total_markets,
                    (SELECT COUNT(*) FROM districts_risk)             AS total_districts,
                    (SELECT COUNT(*) FROM spikes)                     AS total_spikes,
                    (SELECT COUNT(*) FROM spikes WHERE spike_severity='Critical') AS critical,
                    (SELECT COUNT(*) FROM spikes WHERE spike_severity='Severe')   AS severe,
                    (SELECT COUNT(*) FROM spikes WHERE spike_severity='Moderate') AS moderate
            """)
            top_district = await conn.fetchrow("""
                SELECT name_1, region, risk_score FROM districts_risk
                ORDER BY risk_score DESC NULLS LAST LIMIT 1
            """)
            top_commodity = await conn.fetchrow("""
                SELECT commodity, COUNT(*) AS critical_events FROM spikes
                WHERE spike_severity = 'Critical'
                GROUP BY commodity ORDER BY critical_events DESC LIMIT 1
            """)
            districts = await conn.fetch("""
                SELECT name_1, region, risk_score, critical_count, severe_count,
                       moderate_count, spike_rate_pct
                FROM districts_risk ORDER BY risk_score DESC NULLS LAST
            """)
            crit_spikes = await conn.fetch("""
                SELECT date::date::text AS date, district, market, commodity,
                       price, pct_change, zscore
                FROM spikes WHERE spike_severity = 'Critical'
                ORDER BY date DESC LIMIT 30
            """)
            commodities = await conn.fetch("""
                SELECT commodity,
                    COUNT(*) FILTER (WHERE spike_severity='Critical') AS critical_count,
                    COUNT(*) AS total_spikes,
                    ROUND(
                        COUNT(*) FILTER (WHERE spike_severity='Critical')::numeric
                        / NULLIF(COUNT(*),0) * 100, 1
                    ) AS spike_rate
                FROM spikes GROUP BY commodity ORDER BY critical_count DESC LIMIT 20
            """)
            gaps = await conn.fetch("""
                SELECT dr.name_1 AS district, dr.region, dr.risk_score,
                    COUNT(m.ogc_fid) AS market_count,
                    CASE
                        WHEN COUNT(m.ogc_fid) = 0       THEN 'NO MONITORING'
                        WHEN dr.risk_score > 300 AND COUNT(m.ogc_fid) < 3 THEN 'CRITICAL GAP'
                        WHEN dr.risk_score > 150 AND COUNT(m.ogc_fid) < 5 THEN 'HIGH GAP'
                        WHEN dr.risk_score > 50  AND COUNT(m.ogc_fid) < 8 THEN 'MODERATE GAP'
                        ELSE 'ADEQUATE'
                    END AS monitoring_status
                FROM districts_risk dr
                LEFT JOIN markets m ON ST_Within(m.wkb_geometry, dr.wkb_geometry)
                GROUP BY dr.name_1, dr.region, dr.risk_score
                ORDER BY dr.risk_score DESC
            """)

        data = {
            "summary": {
                "total_markets"    : int(summary["total_markets"]),
                "total_districts"  : int(summary["total_districts"]),
                "total_regions"    : "3",
                "spike_events"     : {
                    "total": int(summary["total_spikes"]), "critical": int(summary["critical"]),
                    "severe": int(summary["severe"]), "moderate": int(summary["moderate"]),
                },
                "highest_risk_district": {
                    "name": top_district["name_1"], "region": top_district["region"],
                    "risk_score": float(top_district["risk_score"]),
                },
                "most_spiked_commodity": {
                    "name": top_commodity["commodity"] if top_commodity else "N/A",
                    "critical_events": int(top_commodity["critical_events"]) if top_commodity else 0,
                },
                "monitoring_gaps": {
                    "critical_gap_districts": sum(
                        1 for g in gaps if g["monitoring_status"] in ("CRITICAL GAP","NO MONITORING")
                    ),
                },
            },
            "districts"    : [dict(r) for r in districts],
            "critical_spikes": [dict(r) for r in crit_spikes],
            "commodities"  : [dict(r) for r in commodities],
            "gaps"         : [dict(r) for r in gaps],
        }

        pdf_bytes = build_pdf(data)
        filename  = f"malawi_food_security_{datetime.now().strftime('%Y-%m-%d')}.pdf"
        return Response(
            content=pdf_bytes, media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        import traceback
        return JSONResponse(status_code=500, content={"error": str(e), "trace": traceback.format_exc()})


@router.get("/district/{district_name}")
async def generate_district_report(
    district_name : str,
    date_from     : Optional[str]  = Query(None, description="YYYY-MM-DD"),
    date_to       : Optional[str]  = Query(None, description="YYYY-MM-DD"),
    year          : Optional[str]  = Query(None, description="e.g. 2026"),
    commodity     : List[str]      = Query(default=[], description="One or more commodities"),
    sections      : List[str]      = Query(default=["narrative","prices","spikes","indicators","markets","forecast"]),
    fmt           : str            = Query(default="pdf", description="pdf | docx"),
):
    """
    District-level report with optional date, multi-commodity, and format filters.

    Multi-commodity: pass commodity= multiple times:
      /district/Machinga?commodity=Maize&commodity=Beans
    """
    from urllib.parse import unquote
    district_name = unquote(district_name)

    if year and not date_from:
        date_from = f"{year}-01-01"; date_to = f"{year}-12-31"
    if not date_from: date_from = "2020-01-01"
    if not date_to:   date_to   = datetime.now().strftime("%Y-%m-%d")

    date_from_obj = date_type.fromisoformat(date_from)
    date_to_obj   = date_type.fromisoformat(date_to)

    # Commodity filter: None means all
    comm_list  = [c for c in commodity if c and c != "All commodities"]
    comm_filter = comm_list if comm_list else None  # None = no filter

    try:
        pool = await get_pool()
        async with pool.acquire() as conn:
            district = await conn.fetchrow("""
                SELECT district, region, risk_score, critical_count, severe_count,
                       moderate_count, COALESCE(spike_rate_pct,0) AS spike_rate_pct,
                       COALESCE(avg_price,0) AS avg_price
                FROM districts_risk WHERE district ILIKE $1 LIMIT 1
            """, district_name)
            if not district:
                return JSONResponse(status_code=404,
                                    content={"error": f"District '{district_name}' not found"})

            # spikes — filter by date and (optionally) commodity
            spikes = await conn.fetch("""
                SELECT date::date::text AS date, market, commodity,
                       price, pct_change, zscore, spike_severity
                FROM spikes
                WHERE district ILIKE $1 AND date BETWEEN $2 AND $3
                  AND ($4::text[] IS NULL OR commodity = ANY($4::text[]))
                ORDER BY date DESC LIMIT 100
            """, district_name, date_from_obj, date_to_obj, comm_filter)

            prices = await conn.fetch("""
                SELECT commodity,
                    ROUND(AVG(price::numeric)::numeric, 0)      AS avg_price,
                    ROUND(AVG(pct_change::numeric)::numeric, 1) AS avg_pct_change,
                    COUNT(*)                                     AS obs
                FROM prices
                WHERE district ILIKE $1
                  AND price      ~ '^[0-9]+(\.[0-9]+)?$'
                  AND pct_change ~ '^-?[0-9]+(\.[0-9]+)?$'
                  AND ($2::text[] IS NULL OR commodity = ANY($2::text[]))
                  AND date::date BETWEEN $3 AND $4
                GROUP BY commodity ORDER BY avg_price DESC
            """, district_name, comm_filter, date_from_obj, date_to_obj)

            markets = await conn.fetch("""
                SELECT market, num_commodities, total_spikes, critical_count,
                       spike_rate_pct, latest_date::date::text AS latest_date
                FROM markets WHERE district ILIKE $1
                ORDER BY total_spikes DESC
            """, district_name)

            indicators = await conn.fetchrow("""
                SELECT risk_score, critical_count,
                    COALESCE(spike_rate_pct,0) AS spike_rate_pct,
                    COALESCE(avg_price,0)      AS avg_price,
                    CASE
                        WHEN risk_score > 300 THEN 'CRITICAL GAP'
                        WHEN risk_score > 150 THEN 'HIGH GAP'
                        WHEN risk_score > 50  THEN 'MODERATE GAP'
                        ELSE 'ADEQUATE'
                    END AS monitoring_status
                FROM districts_risk WHERE district ILIKE $1
            """, district_name)

            # forecast data from forecasts table (if it exists)
            forecast = None
            try:
                fc_rows = await conn.fetch("""
                    SELECT month_label, season, predicted_price,
                           pct_vs_baseline, severity
                    FROM forecasts
                    WHERE district ILIKE $1
                      AND ($2::text[] IS NULL OR commodity = ANY($2::text[]))
                    ORDER BY forecast_date ASC LIMIT 3
                """, district_name, comm_filter)
                if fc_rows:
                    forecast = {"months": [dict(r) for r in fc_rows]}
            except Exception:
                pass  # forecasts table may not exist yet

        payload = {
            "district"     : dict(district),
            "spikes"       : [dict(r) for r in spikes],
            "prices"       : [dict(r) for r in prices],
            "markets"      : [dict(r) for r in markets],
            "sections"     : sections,
            "date_range"   : {"from": date_from, "to": date_to},
            "commodity_list": comm_list or ["All commodities"],
            "indicators"   : dict(indicators) if indicators else None,
            "forecast"     : forecast,
        }

        safe_name = district_name.replace(" ", "_").lower()
        date_tag  = datetime.now().strftime('%Y-%m-%d')

        if fmt.lower() == "docx":
            docx_bytes = build_district_docx(payload)
            filename   = f"report_{safe_name}_{date_tag}.docx"
            return Response(
                content=docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={filename}"},
            )

        pdf_bytes = build_district_pdf(payload)
        filename  = f"report_{safe_name}_{date_tag}.pdf"
        return Response(
            content=pdf_bytes, media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except Exception as e:
        import traceback
        return JSONResponse(status_code=500,
                            content={"error": str(e), "trace": traceback.format_exc()})